# export_dialog.py
from __future__ import annotations

import os
from dataclasses import asdict
from datetime import datetime
from decimal import Decimal
from typing import List, Dict, Any, Optional, Callable
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from PyQt6.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, QPushButton,
    QComboBox, QProgressBar, QMessageBox, QFileDialog
)

from am_models import Ammunition, SQLRepository

# ---------------------- 工具函数：序列化 Ammunition 为字典 ----------------------
_AM_FIELDS = [
    # 主键
    "am_id",
    # 基本必填
    "am_name","chinese_name", "short_name", "am_type", "weight_kg", "launch_mass_kg",
    "warhead_type", "warhead_name",
    # 可选
    "country", "used_by",  "model_name",
    "submodel_name", "manufacturer", "attended_time",
    "length_m", "diameter_m", "texture", "wingspan_close_mm",
    "wingspan_open_mm", "structure",
    "max_speed_ma", "radar_cross_section", "power_plant",
    "destroying_elements", "fuze", "explosion_equivalent_TNT_T",
    "precision_m", "destroying_mechanism",
    "target", "carrier", "guidance_mode", "explosive_payload_kg",
    "penetrating_power", "drop_height_range_m", "drop_speed_kmh",
    "drop_mode", "coverage_area", "range_km",

    # 爆破战斗部
    "is_explosive_bomb",
    "exb_component",
    "exb_explosion",
    "exb_weight",
    "exb_more_parameters",

    # 聚能战斗部
    "is_energy_bomb",
    "eb_density",
    "eb_velocity",
    "eb_pressure",
    "eb_cover_material",
    "eb_cone_angle",
    "eb_more_parameters",

    # 破片战斗部
    "is_fragment_bomb",
    "fb_bomb_explosion",
    "fb_fragment_shape",
    "fb_surface_area",
    "fb_fragment_weight",
    "fb_diameter",
    "fb_length",
    "fb_shell_weight",
    "fb_more_parameters",

    # 穿甲战斗部
    "is_armor_bomb",
    "ab_bullet_weight",
    "ab_diameter",
    "ab_head_length",
    "ab_more_parameters",

    # 子母弹
    "is_cluster_bomb",
    "cbm_bullet_weight",
    "cbm_bullet_section",
    "cbm_projectile",
    "cbs_bullet_count",
    "cbs_bullet_model",
    "cbs_bullet_weight",
    "cb_diameter",
    "cbs_bullet_length",
    "cb_more_parameters",

    # 审计
    "created_at", "updated_at",
]

FIELD_ORDER = [
    "am_id",
    "am_name",
    "chinese_name",
    "short_name",
    "country",
    "used_by",
    "am_type",
    "model_name",
    "submodel_name",
    "manufacturer",
    "attended_time",
    "weight_kg",
    "length_m",
    "diameter_m",
    "texture",
    "wingspan_close_mm",
    "wingspan_open_mm",
    "structure",
    "max_speed_ma",
    "radar_cross_section",
    "power_plant",
    "launch_mass_kg",
    "warhead_type",
    "warhead_name",
    "destroying_elements",
    "fuze",
    "explosion_equivalent_TNT_T",
    "precision_m",
    "destroying_mechanism",
    "target",
    "carrier",
    "guidance_mode",
    "explosive_payload_kg",
    "penetrating_power",
    "drop_height_range_m",
    "drop_speed_kmh",
    "drop_mode",
    "coverage_area",
    "range_km",

    # EXB
    "is_explosive_bomb",
    "exb_component",
    "exb_explosion",
    "exb_weight",
    "exb_more_parameters",

    # EB
    "is_energy_bomb",
    "eb_density",
    "eb_velocity",
    "eb_pressure",
    "eb_cover_material",
    "eb_cone_angle",
    "eb_more_parameters",

    # FB
    "is_fragment_bomb",
    "fb_bomb_explosion",
    "fb_fragment_shape",
    "fb_surface_area",
    "fb_fragment_weight",
    "fb_diameter",
    "fb_length",
    "fb_shell_weight",
    "fb_more_parameters",

    # AB
    "is_armor_bomb",
    "ab_bullet_weight",
    "ab_diameter",
    "ab_head_length",
    "ab_more_parameters",

    # CB/CBS
    "is_cluster_bomb",
    "cbm_bullet_weight",
    "cbm_bullet_section",
    "cbm_projectile",
    "cbs_bullet_count",
    "cbs_bullet_model",
    "cbs_bullet_weight",
    "cb_diameter",
    "cbs_bullet_length",
    "cb_more_parameters",

    "created_at",
    "updated_at",
]

# —— 中文表头（与上面 FIELD_ORDER 完全对应）——
CH_HEADERS = [
    "自增主键", "官方名称", "中文名称", "简称", "国家", "使用单位", "弹药类型",
    "弹药型号", "型号子类", "制造商", "服役时间", "弹体全重", "弹体长度", "弹体直径",
    "弹体材质", "翼展(闭合)", "翼展(张开)", "结构", "最大时速", "雷达截面", "动力装置",
    "发射质量", "战斗部类型", "战斗部", "毁伤元", "引信", "爆炸当量(TNT)", "精度(圆概率误差CEP)",
    "破坏机制", "打击目标", "载机(投放平台)", "制导方式", "装药量", "穿透能力", "投弹高度范围",
    "投弹速度", "投弹方式", "布撒范围", "射程",

    # EXB
    "是否爆破战斗部", "炸药成分", "热爆(爆热)", "装药质量", "爆破其他参数",

    # EB
    "是否聚能战斗部", "炸药密度", "爆速", "爆轰压", "覆盖材料", "锥角", "聚能其他参数",

    # FB
    "是否破片战斗部", "热爆", "破片形状", "破片表面积", "破片质量",
    "装药直径", "装药长度", "壳体质量", "破片其他参数",

    # AB
    "是否穿甲战斗部", "弹丸质量", "弹丸直径", "弹丸头部长度", "穿甲其他参数",

    # CB/CBS
    "是否子母弹", "母弹质量", "母弹最大横截面", "母弹阻力系数",
    "子弹数量", "子弹型号", "子弹质量", "最大直径", "子弹参考长度", "子母弹其他参数",

    "创建时间(UTC)", "更新时间(UTC)"
]


def _to_str(x: Any) -> str:
    if x is None:
        return ""
    if isinstance(x, Decimal):
        return str(x)  # 避免二进制浮点误差
    if isinstance(x, datetime):
        return x.strftime("%Y-%m-%d %H:%M:%S")
    return str(x)


def ammo_to_row_dict(a: Ammunition) -> Dict[str, str]:
    d = {}
    for k in _AM_FIELDS:
        v = getattr(a, k, None)
        d[k] = _to_str(v)
    return d


# ---------------------- 后台线程：执行导出 ----------------------
class ExportWorker(QThread):
    progress = pyqtSignal(int)  # 0..100
    message = pyqtSignal(str)  # 状态文本
    error = pyqtSignal(str)  # 错误
    done = pyqtSignal(str)  # 完成，携带文件路径

    def __init__(
            self,
            out_dir: str,
            fmt: str,
            session_scope: Optional[Callable] = None,
            items: Optional[List[Ammunition]] = None,
            parent=None,
    ):
        """
        两种来源二选一：
          - items 不为 None：直接导出这些记录
          - 否则使用 session_scope() 从数据库读取全部记录
        """
        super().__init__(parent)
        self.out_dir = out_dir
        self.fmt = fmt.lower().strip()
        self.session_scope = session_scope
        self.items = items

    def run(self):
        # 1) 获取待导出的数据
        try:
            self.message.emit("正在读取数据 ...")
            if self.items is not None:
                items: List[Ammunition] = list(self.items)
            else:
                if self.session_scope is None:
                    self.error.emit("未提供数据来源：缺少 items 或 session_scope")
                    return
                with self.session_scope() as s:
                    repo = SQLRepository(s)
                    items = repo.list_all()
        except Exception as e:
            self.error.emit(f"读取数据失败：{e}")
            return

        total = max(len(items), 1)
        self.progress.emit(5)

        # 2) 生成目标文件名
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        if self.fmt == "excel":
            filename = os.path.join(self.out_dir, f"Ammunitions_{ts}.xlsx")
        elif self.fmt == "word":
            filename = os.path.join(self.out_dir, f"Ammunitions_{ts}.docx")
        else:
            self.error.emit(f"未知导出格式：{self.fmt}")
            return

        # 3) 构建行数据并发进度
        self.message.emit("正在准备数据 ...")
        rows: List[Dict[str, str]] = []
        for i, a in enumerate(items, start=1):
            rows.append(ammo_to_row_dict(a))
            # 读数阶段推进到 70%
            pct = 5 + int(60 * i / total)
            self.progress.emit(min(pct, 70))

        # 4) 写文件
        try:
            if self.fmt == "excel":
                self._write_excel(filename, rows)
            else:
                self._write_word(filename, rows)
        except ImportError as e:
            self.error.emit(
                f"缺少导出依赖：{e}. \nExcel 需 pandas+openpyxl，Word 需 python-docx。"
            )
            return
        except Exception as e:
            self.error.emit(f"写文件失败：{e}")
            return

        self.progress.emit(100)
        self.message.emit("导出完成")
        self.done.emit(filename)

    # --- 写 Excel ---
    def _write_excel(self, filename: str, rows: List[Dict[str, str]]):
        """
        使用 pandas+openpyxl 写入，并用 openpyxl 调整样式：首行中文、加粗、底色、细边框、自动列宽、冻结首行
        """
        self.message.emit("正在写入 Excel ...")
        import pandas as pd  # type: ignore
        from openpyxl.styles import Font, PatternFill, Border, Side, Alignment  # type: ignore
        from openpyxl.utils import get_column_letter  # type: ignore

        # 1) 先把数据按 FIELD_ORDER 排序，并映射为中文表头
        ordered_rows = []
        for r in rows:
            ordered_rows.append({zh: r.get(k, "") for k, zh in zip(FIELD_ORDER, CH_HEADERS)})

        df = pd.DataFrame(ordered_rows, columns=CH_HEADERS)

        # 2) 写入 Excel
        with pd.ExcelWriter(filename, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="Ammunition")
            wb = writer.book
            ws = writer.sheets["Ammunition"]

            # 3) 样式：首行加粗、底色、居中；所有单元格加细边框；冻结首行
            header_font = Font(bold=True)
            header_fill = PatternFill("solid", fgColor="DDDDDD")
            thin = Side(border_style="thin", color="888888")
            border = Border(top=thin, bottom=thin, left=thin, right=thin)
            center = Alignment(vertical="center")

            max_col = ws.max_column
            max_row = ws.max_row

            # 首行样式
            for col in range(1, max_col + 1):
                cell = ws.cell(row=1, column=col)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = center
                cell.border = border

            # 数据区样式 + 简单“自动列宽”
            col_widths = [len(CH_HEADERS[c - 1]) + 5 for c in range(1, max_col + 1)]
            for r in range(2, max_row + 1):
                for c in range(1, max_col + 1):
                    cell = ws.cell(row=r, column=c)
                    cell.alignment = center
                    cell.border = border
                    val = cell.value
                    if val is None:
                        l = 0
                    else:
                        l = len(str(val))
                    if l + 2 > col_widths[c - 1]:
                        col_widths[c - 1] = min(l + 2, 60)  # 上限 60，避免过宽

            for c, w in enumerate(col_widths, start=1):
                ws.column_dimensions[get_column_letter(c)].width = w

            # 冻结首行
            ws.freeze_panes = "A2"

    # --- 写 Word ---
    def _write_word(self, filename: str, rows: list[dict]):
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.shared import Pt

        doc = Document()

        # 设置文档默认字体：中文=宋体，英文字体=Times New Roman
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        # eastAsia 中文字体设置
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # 标题
        title = doc.add_heading("Ammunition 导出", level=1)
        # 为标题也设置中英文字体
        title.runs[0].font.name = "Times New Roman"
        title.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # 每条数据：编号．名称
        total = max(len(rows), 1)
        for i, r in enumerate(rows, start=1):
            # 进度（70% -> 100%）
            pct = 70 + int(30 * i / total)
            self.progress.emit(min(pct, 99))

            am_id = r.get("am_id", "")
            am_name = r.get("am_name", "")

            # 标题行：编号．名称
            title_p = doc.add_paragraph(f"{am_id}．{am_name}")
            title_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            # 强制中英文字体
            for run in title_p.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(12)

            # 符号列表内容：除“自增主键/名称”外的所有字段，逐条 "字段名：值"
            for header, field in zip(CH_HEADERS, FIELD_ORDER):
                if field in ("am_id", "am_name"):  # 已放在标题中
                    continue
                val = r.get(field, "")
                p = doc.add_paragraph(style=doc.styles["List Bullet"])
                run = p.add_run(f"{header}：{val}")
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(11)

        doc.save(filename)


# ---------------------- 导出对话框（UI + 逻辑）----------------------
class ExportDialog(QDialog):
    def __init__(
            self,
            parent,
            session_scope: Optional[Callable] = None,
            items: Optional[List[Ammunition]] = None,
    ):
        """
        - 仅传 session_scope：导出数据库全部记录
        - 仅传 items：导出这批记录
        - 两者都传：优先 items（导出检索结果等）
        """
        super().__init__(parent)
        self.setWindowTitle("导出弹药毁伤数据模型")
        self.setModal(True)
        self.resize(500, 180)

        self.session_scope = session_scope
        self.items = items
        self.worker: Optional[ExportWorker] = None

        vbox = QVBoxLayout(self)

        # 行1：目录 + 浏览
        row1 = QHBoxLayout()
        row1.addWidget(QLabel("导出目录："))
        self.ed_dir = QLineEdit()
        self.btn_browse = QPushButton("浏览...")
        self.btn_browse.clicked.connect(self._on_browse)
        row1.addWidget(self.ed_dir, 1)
        row1.addWidget(self.btn_browse)
        vbox.addLayout(row1)

        # 行2：格式
        row2 = QHBoxLayout()
        row2.addWidget(QLabel("导出格式："))
        self.cmb_fmt = QComboBox()
        self.cmb_fmt.addItems(["Excel", "Word"])
        row2.addWidget(self.cmb_fmt, 1)
        vbox.addLayout(row2)

        # 行3：按钮
        row3 = QHBoxLayout()
        self.btn_export = QPushButton("开始导出")
        self.btn_export.clicked.connect(self._on_export)
        row3.addStretch(1)
        row3.addWidget(self.btn_export)
        vbox.addLayout(row3)

        # 进度条 + 状态
        self.progress = QProgressBar()
        self.progress.setRange(0, 100)
        self.progress.setValue(0)
        vbox.addWidget(self.progress)

        self.lbl_status = QLabel("")
        self.lbl_status.setStyleSheet("color: gray;")
        vbox.addWidget(self.lbl_status)

    # 选择目录
    def _on_browse(self):
        path = QFileDialog.getExistingDirectory(self, "选择导出目录", os.path.expanduser("~"))
        if path:
            self.ed_dir.setText(path)

    # 开始导出
    def _on_export(self):
        out_dir = self.ed_dir.text().strip()
        if not out_dir:
            QMessageBox.warning(self, "提示", "请先选择导出目录")
            return
        if not os.path.isdir(out_dir):
            QMessageBox.warning(self, "提示", "导出目录无效")
            return

        fmt = self.cmb_fmt.currentText().strip().lower()
        self.progress.setValue(0)
        self.lbl_status.setText("正在启动导出 ...")
        self.btn_export.setEnabled(False)

        # 后台线程：若 items 非空，优先用 items；否则用 session_scope 全量导出
        self.worker = ExportWorker(
            out_dir=out_dir,
            fmt=fmt,
            session_scope=self.session_scope,
            items=self.items,
            parent=self,
        )
        self.worker.progress.connect(self.progress.setValue)
        self.worker.message.connect(self.lbl_status.setText)
        self.worker.error.connect(self._on_error)
        self.worker.done.connect(self._on_done)
        self.worker.finished.connect(lambda: self.btn_export.setEnabled(True))
        self.worker.start()

    def _on_error(self, msg: str):
        self.lbl_status.setText(msg)
        QMessageBox.critical(self, "导出失败", msg)

    def _on_done(self, filename: str):
        self.lbl_status.setText(f"导出完成：{filename}")
        QMessageBox.information(self, "导出完成", f"已导出到：\n{filename}")
        # self.accept()


# ---------------------- 便捷入口 ----------------------
def show_export_dialog(parent, session_scope: Optional[Callable] = None, items: Optional[List[Ammunition]] = None):
    """
    一个简单的 helper：任何地方都可调用
        - 导出检索结果： show_export_dialog(self, items=found_list)
        - 全量导出：     show_export_dialog(self, session_scope=self.session_scope)
    """
    dlg = ExportDialog(parent, session_scope=session_scope, items=items)
    dlg.exec()
