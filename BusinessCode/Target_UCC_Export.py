from __future__ import annotations

import base64
import os
from datetime import datetime
from decimal import Decimal
from typing import Any, Dict, List, Optional, Sequence

from PyQt6.QtCore import QThread, pyqtSignal
from PyQt6.QtWidgets import (
    QComboBox,
    QDialog,
    QFileDialog,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QMessageBox,
    QPushButton,
    QProgressBar,
    QVBoxLayout,
)


UG_FIELD_ORDER: Sequence[str] = (
    "id",
    "ucc_code",
    "ucc_name",
    "country",
    "base",
    "location",
    "shelter_picture",
    "rock_layer_materials",
    "rock_layer_thick",
    "rock_layer_strength",
    "protective_layer_material",
    "protective_layer_thick",
    "protective_layer_strength",
    "lining_layer_material",
    "lining_layer_thick",
    "lining_layer_strength",
    "ucc_wall_materials",
    "ucc_wall_thick",
    "ucc_wall_strength",
    "ucc_length",
    "ucc_width",
    "ucc_height",
    "ucc_status",
    "created_time",
    "updated_time",
)


UG_HEADERS: Sequence[str] = (
    "ID",
    "\u6307\u6325\u6240\u4ee3\u7801",
    "\u6307\u6325\u6240\u540d\u79f0",
    "\u56fd\u5bb6/\u5730\u533a",
    "\u57fa\u5730/\u90e8\u961f",
    "\u6240\u5728\u4f4d\u7f6e",
    "\u7167\u7247\u8def\u5f84",
    "\u571f\u58e4\u5ca9\u5c42\u6750\u6599",
    "\u571f\u58e4\u5ca9\u5c42\u539a\u5ea6(cm)",
    "\u571f\u58e4\u5ca9\u5c42\u6297\u538b\u5f3a\u5ea6(MPa)",
    "\u9632\u62a4\u5c42\u6750\u6599",
    "\u9632\u62a4\u5c42\u539a\u5ea6(cm)",
    "\u9632\u62a4\u5c42\u6297\u538b\u5f3a\u5ea6(MPa)",
    "\u8863\u91c7\u5c42\u6750\u6599",
    "\u8863\u91c7\u5c42\u539a\u5ea6(cm)",
    "\u8863\u91c7\u5c42\u6297\u538b\u5f3a\u5ea6(MPa)",
    "UCC\u5899\u4f53\u6750\u6599",
    "UCC\u5899\u4f53\u539a\u5ea6(cm)",
    "UCC\u5899\u4f53\u6297\u538b\u5f3a\u5ea6(MPa)",
    "\u7a7a\u95f4\u957f\u5ea6(m)",
    "\u7a7a\u95f4\u5bbd\u5ea6(m)",
    "\u7a7a\u95f4\u9ad8\u5ea6(m)",
    "\u6307\u6325\u6240\u72b6\u6001",
    "\u521b\u5efa\u65f6\u95f4(UTC)",
    "\u66f4\u65b0\u65f6\u95f4(UTC)",
)




def _to_str(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, Decimal):
        return str(value)
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d %H:%M:%S")
    if isinstance(value, (bytes, bytearray)):
        try:
            return base64.b64encode(value).decode("ascii")
        except Exception:
            return "[binary]"
    if isinstance(value, bool):
        return "\u662f" if value else "\u5426"
    if isinstance(value, (int, float)):
        formatted = f"{float(value):.4f}".rstrip("0").rstrip(".")
        return formatted
    return str(value)



def post_to_dict(obj: Any) -> Dict[str, str]:
    return {field: _to_str(getattr(obj, field, None)) for field in UG_FIELD_ORDER}


class ExportWorker(QThread):
    progress = pyqtSignal(int)
    message = pyqtSignal(str)
    error = pyqtSignal(str)
    done = pyqtSignal(str)

    def __init__(self, session_scope, out_dir: str, fmt: str, parent=None):
        super().__init__(parent)
        self.session_scope = session_scope
        self.out_dir = out_dir
        self.fmt = fmt.lower().strip()

    def run(self) -> None:  # pragma: no cover - GUI thread
        try:
            from target_model.sql_repository import SQLRepository
            from target_model.entities import UndergroundCommandPost
        except Exception as exc:
            self.error.emit(f"导出依赖加载失败：{exc}")
            return

        try:
            self.message.emit("正在读取地下指挥所数据 ...")
            with self.session_scope() as session:
                repo = SQLRepository(session)
                posts = repo.list_all(UndergroundCommandPost)
        except Exception as exc:
            self.error.emit(f"读取数据失败：{exc}")
            return

        total = max(len(posts), 1)
        self.progress.emit(10)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        if self.fmt == "excel":
            output = os.path.join(self.out_dir, f"UndergroundPosts_{timestamp}.xlsx")
        elif self.fmt == "word":
            output = os.path.join(self.out_dir, f"UndergroundPosts_{timestamp}.docx")
        else:
            self.error.emit(f"未知导出格式：{self.fmt}")
            return

        self.message.emit("正在整理数据 ...")
        rows: List[Dict[str, str]] = []
        for index, post in enumerate(posts, start=1):
            rows.append(post_to_dict(post))
            pct = 10 + int(60 * index / total)
            self.progress.emit(min(pct, 70))

        try:
            if self.fmt == "excel":
                self._write_excel(output, rows)
            else:
                self._write_word(output, rows)
        except ImportError as exc:
            self.error.emit(
                f"缺少导出依赖：{exc}。Excel 导出需要 pandas+openpyxl，Word 导出需要 python-docx。"
            )
            return
        except Exception as exc:
            self.error.emit(f"写文件失败：{exc}")
            return

        self.progress.emit(100)
        self.message.emit("导出完成")
        self.done.emit(output)

    def _write_excel(self, filename: str, rows: List[Dict[str, str]]) -> None:
        self.message.emit("正在写入 Excel ...")
        import pandas as pd  # type: ignore
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side  # type: ignore
        from openpyxl.utils import get_column_letter  # type: ignore

        ordered = [
            {header: row.get(field, "") for field, header in zip(UG_FIELD_ORDER, UG_HEADERS)}
            for row in rows
        ]
        df = pd.DataFrame(ordered, columns=UG_HEADERS)

        with pd.ExcelWriter(filename, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="UndergroundPosts")
            ws = writer.sheets["UndergroundPosts"]

            header_font = Font(bold=True)
            header_fill = PatternFill("solid", fgColor="DDDDDD")
            thin = Side(border_style="thin", color="999999")
            border = Border(top=thin, bottom=thin, left=thin, right=thin)
            center = Alignment(vertical="center")

            for row in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=ws.max_column):
                for cell in row:
                    cell.border = border
                    if cell.row == 1:
                        cell.font = header_font
                        cell.fill = header_fill
                        cell.alignment = center

            for col_idx in range(1, ws.max_column + 1):
                column = get_column_letter(col_idx)
                max_length = max(
                    (len(str(ws.cell(row=row, column=col_idx).value or "")) for row in range(1, ws.max_row + 1)),
                    default=0,
                )
                ws.column_dimensions[column].width = min(max(max_length + 2, 12), 60)

            ws.freeze_panes = "A2"

    def _write_word(self, filename: str, rows: List[Dict[str, str]]) -> None:
        self.message.emit("正在写入 Word ...")
        from docx import Document  # type: ignore
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.shared import Pt  # type: ignore

        doc = Document()
        doc.add_heading("地下指挥所数据", level=1)


        for row in rows or [{}]:
            code = row.get("ucc_code", "")
            name = row.get("ucc_name", "")
            post_id = row.get("id", "")
            if code and name:
                display = f"{code} · {name}"
            else:
                display = code or name or "未命名指挥所"
            heading_parts = [part for part in (str(post_id).strip(), display) if part]
            title_para = doc.add_paragraph(" · ".join(heading_parts))
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in title_para.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(12)

            for header, field in zip(UG_HEADERS, UG_FIELD_ORDER):
                if field in ("id", "ucc_code", "ucc_name"):
                    continue
                value = row.get(field, "")
                paragraph = doc.add_paragraph(style="List Bullet")
                run = paragraph.add_run(f"{header}: {value}")
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(11)

        doc.save(filename)


class Target_UCC_ExportWindow(QDialog):
    def __init__(self, parent, session_scope) -> None:
        super().__init__(parent)
        self.setWindowTitle("导出地下指挥所数据模型")
        self.setModal(True)
        self.resize(500, 180)

        self.session_scope = session_scope
        self.worker: Optional[ExportWorker] = None

        layout = QVBoxLayout(self)

        row1 = QHBoxLayout()
        row1.addWidget(QLabel("导出目录："))
        self.ed_dir = QLineEdit()
        self.btn_browse = QPushButton("浏览...")
        self.btn_browse.clicked.connect(self._on_browse)
        row1.addWidget(self.ed_dir, 1)
        row1.addWidget(self.btn_browse)
        layout.addLayout(row1)

        row2 = QHBoxLayout()
        row2.addWidget(QLabel("导出格式："))
        self.cmb_fmt = QComboBox()
        self.cmb_fmt.addItems(["Excel", "Word"])
        row2.addWidget(self.cmb_fmt, 1)
        layout.addLayout(row2)

        row3 = QHBoxLayout()
        row3.addStretch(1)
        self.btn_export = QPushButton("开始导出")
        self.btn_export.clicked.connect(self._on_export)
        row3.addWidget(self.btn_export)
        layout.addLayout(row3)

        self.progress = QProgressBar()
        self.progress.setRange(0, 100)
        layout.addWidget(self.progress)

        self.lbl_status = QLabel("")
        self.lbl_status.setStyleSheet("color: gray;")
        layout.addWidget(self.lbl_status)

    def _on_browse(self) -> None:
        directory = QFileDialog.getExistingDirectory(self, "选择导出目录", os.path.expanduser("~"))
        if directory:
            self.ed_dir.setText(directory)

    def _on_export(self) -> None:
        out_dir = self.ed_dir.text().strip()
        if not out_dir:
            QMessageBox.warning(self, "提示", "请先选择导出目录")
            return
        if not os.path.isdir(out_dir):
            QMessageBox.warning(self, "提示", "导出目录无效")
            return

        fmt = self.cmb_fmt.currentText().strip().lower()

        self.progress.setValue(0)
        self.lbl_status.setText("正在启动导出 ...")
        self.btn_export.setEnabled(False)

        self.worker = ExportWorker(self.session_scope, out_dir, fmt, self)
        self.worker.progress.connect(self.progress.setValue)
        self.worker.message.connect(self.lbl_status.setText)
        self.worker.error.connect(self._on_error)
        self.worker.done.connect(self._on_done)
        self.worker.finished.connect(lambda: self.btn_export.setEnabled(True))
        self.worker.start()

    def _on_error(self, message: str) -> None:
        self.lbl_status.setText(message)
        QMessageBox.critical(self, "导出失败", message)

    def _on_done(self, filename: str) -> None:
        self.lbl_status.setText(f"导出完成：{filename}")
        QMessageBox.information(self, "导出完成", f"已导出到：\n{filename}")

