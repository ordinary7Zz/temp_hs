from __future__ import annotations

import base64
import os
from datetime import datetime
from decimal import Decimal
from typing import Any, Dict, List, Optional, Sequence

from PyQt6.QtCore import QThread, pyqtSignal
from PyQt6.QtWidgets import (
    QComboBox,
    QDialog,
    QFileDialog,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QMessageBox,
    QPushButton,
    QProgressBar,
    QVBoxLayout,
)

RUNWAY_FIELD_ORDER: Sequence[str] = (
    "id",
    "runway_code",
    "runway_name",
    "country",
    "base",
    "runway_picture",
    "r_length",
    "r_width",
    "pccsc_thick",
    "pccsc_strength",
    "pccsc_flexural",
    "pccsc_freeze",
    "pccsc_cement",
    "pccsc_block_size1",
    "pccsc_block_size2",
    "ctbc_thick",
    "ctbc_strength",
    "ctbc_flexural",
    "ctbc_cement",
    "ctbc_compaction",
    "gcss_thick",
    "gcss_strength",
    "gcss_compaction",
    "cs_thick",
    "cs_strength",
    "cs_compaction",
    "runway_status",
    "created_time",
    "updated_time",
)

RUNWAY_HEADERS_ZH: Sequence[str] = (
    "编号",
    "跑道代码",
    "机场名称",
    "国家/地区",
    "基地/部队",
    "机场照片",
    "跑道长度(m)",
    "跑道宽度(m)",
    "混凝土面层厚度(cm)",
    "混凝土面层抗压强度(MPa)",
    "混凝土面层抗折强度(MPa)",
    "抗冻融循环次数",
    "水泥类型",
    "道面分块尺寸1(m)",
    "道面分块尺寸2(m)",
    "水泥稳定碎石基层厚度(cm)",
    "水泥稳定碎石基层抗压强度(MPa)",
    "水泥稳定碎石基层抗折强度(MPa)",
    "水泥掺量",
    "夯实密实度",
    "级配砂砾石垫层厚度(cm)",
    "级配砂砾石垫层强度承载比(%)",
    "级配砂砾石垫层压实模量(MPa)",
    "土基压实层厚度(cm)",
    "土基压实层强度承载比(%)",
    "土基压实层压实模量(MPa)",
    "跑道状态",
    "创建时间(UTC)",
    "更新时间(UTC)",
)



class Target_Runway_ExportWindow(QDialog):
    def __init__(self, parent, session_scope) -> None:
        super().__init__(parent)
        self.setWindowTitle("导出机场跑道数据模型")
        self.setModal(True)
        self.resize(500, 180)

        self.session_scope = session_scope
        self.worker: Optional[ExportWorker] = None

        layout = QVBoxLayout(self)

        row1 = QHBoxLayout()
        row1.addWidget(QLabel("导出目录："))
        self.ed_dir = QLineEdit()
        self.btn_browse = QPushButton("浏览...")
        self.btn_browse.clicked.connect(self._on_browse)
        row1.addWidget(self.ed_dir, 1)
        row1.addWidget(self.btn_browse)
        layout.addLayout(row1)

        row2 = QHBoxLayout()
        row2.addWidget(QLabel("导出格式："))
        self.cmb_fmt = QComboBox()
        self.cmb_fmt.addItems(["Excel", "Word"])
        row2.addWidget(self.cmb_fmt, 1)
        layout.addLayout(row2)

        row3 = QHBoxLayout()
        row3.addStretch(1)
        self.btn_export = QPushButton("开始导出")
        self.btn_export.clicked.connect(self._on_export)
        row3.addWidget(self.btn_export)
        layout.addLayout(row3)

        self.progress = QProgressBar()
        self.progress.setRange(0, 100)
        self.progress.setValue(0)
        layout.addWidget(self.progress)

        self.lbl_status = QLabel("")
        self.lbl_status.setStyleSheet("color: gray;")
        layout.addWidget(self.lbl_status)

    def _on_browse(self) -> None:
        directory = QFileDialog.getExistingDirectory(self, "选择导出目录", os.path.expanduser("~"))
        if directory:
            self.ed_dir.setText(directory)

    def _on_export(self) -> None:
        out_dir = self.ed_dir.text().strip()
        if not out_dir:
            QMessageBox.warning(self, "提示", "请先选择导出目录")
            return
        if not os.path.isdir(out_dir):
            QMessageBox.warning(self, "提示", "导出目录无效")
            return

        fmt = self.cmb_fmt.currentText().strip().lower()

        self.progress.setValue(0)
        self.lbl_status.setText("正在启动导出 ...")
        self.btn_export.setEnabled(False)

        self.worker = ExportWorker(self.session_scope, out_dir, fmt, self)
        self.worker.progress.connect(self.progress.setValue)
        self.worker.message.connect(self.lbl_status.setText)
        self.worker.error.connect(self._on_error)
        self.worker.done.connect(self._on_done)
        self.worker.finished.connect(lambda: self.btn_export.setEnabled(True))
        self.worker.start()

    def _on_error(self, message: str) -> None:
        self.lbl_status.setText(message)
        QMessageBox.critical(self, "导出失败", message)

    def _on_done(self, filename: str) -> None:
        self.lbl_status.setText(f"导出完成：{filename}")
        QMessageBox.information(self, "导出完成", f"已导出到：\n{filename}")



class ExportWorker(QThread):
    progress = pyqtSignal(int)
    message = pyqtSignal(str)
    error = pyqtSignal(str)
    done = pyqtSignal(str)

    def __init__(self, session_scope, out_dir: str, fmt: str, parent=None):
        super().__init__(parent)
        self.session_scope = session_scope
        self.out_dir = out_dir
        self.fmt = fmt.lower().strip()

    def run(self) -> None:
        try:
            from target_model.sql_repository import SQLRepository
            from target_model.entities import AirportRunway
        except Exception as exc:  # pragma: no cover - import guard
            self.error.emit(f"导出依赖加载失败：{exc}")
            return

        try:
            self.message.emit("正在读取跑道数据 ...")
            with self.session_scope() as session:
                repo = SQLRepository(session)
                runways = repo.list_all(AirportRunway)
        except Exception as exc:
            self.error.emit(f"读取跑道数据失败：{exc}")
            return

        total = max(len(runways), 1)
        self.progress.emit(10)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        if self.fmt == "excel":
            outfile = os.path.join(self.out_dir, f"Runways_{timestamp}.xlsx")
        elif self.fmt == "word":
            outfile = os.path.join(self.out_dir, f"Runways_{timestamp}.docx")
        else:
            self.error.emit(f"未知导出格式：{self.fmt}")
            return

        self.message.emit("正在整理数据 ...")
        rows = []
        for idx, runway in enumerate(runways, start=1):
            rows.append(runway_to_dict(runway))
            pct = 10 + int(60 * idx / total)
            self.progress.emit(min(pct, 70))

        try:
            if self.fmt == "excel":
                self._write_excel(outfile, rows)
            else:
                self._write_word(outfile, rows)
        except ImportError as exc:
            self.error.emit(
                f"缺少导出依赖：{exc}。Excel 导出需要 pandas+openpyxl，Word 导出需要 python-docx。"
            )
            return
        except Exception as exc:
            self.error.emit(f"写文件失败：{exc}")
            return

        self.progress.emit(100)
        self.message.emit("导出完成")
        self.done.emit(outfile)

    def _write_excel(self, filename: str, rows: List[Dict[str, str]]) -> None:
        self.message.emit("正在写入 Excel ...")
        import pandas as pd  # type: ignore
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side  # type: ignore
        from openpyxl.utils import get_column_letter  # type: ignore

        ordered = [
            {header: row.get(field, "") for field, header in zip(RUNWAY_FIELD_ORDER, RUNWAY_HEADERS_ZH)}
            for row in rows
        ]

        df = pd.DataFrame(ordered, columns=RUNWAY_HEADERS_ZH)
        with pd.ExcelWriter(filename, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="Runways")
            ws = writer.sheets["Runways"]

            header_font = Font(bold=True)
            header_fill = PatternFill("solid", fgColor="DDDDDD")
            thin = Side(border_style="thin", color="999999")
            border = Border(top=thin, bottom=thin, left=thin, right=thin)
            center = Alignment(vertical="center")

            for row in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=ws.max_column):
                for cell in row:
                    cell.border = border
                    if cell.row == 1:
                        cell.font = header_font
                        cell.fill = header_fill
                        cell.alignment = center

            for col_idx in range(1, ws.max_column + 1):
                column = get_column_letter(col_idx)
                max_length = max((len(str(ws.cell(row=row, column=col_idx).value or "")) for row in range(1, ws.max_row + 1)), default=0)
                ws.column_dimensions[column].width = min(max(max_length + 2, 12), 60)

            ws.freeze_panes = "A2"

    def _write_word(self, filename: str, rows: List[Dict[str, str]]) -> None:
        self.message.emit("正在写入 Word ...")
        from docx import Document  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.shared import Pt  # type: ignore
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # type: ignore

        doc = Document()
        doc.add_heading("机场跑道数据模型导出", level=1)

        for row in rows or [{}]:
            title = row.get("runway_name", "") or "未命名跑道"
            run_id = row.get("id", "")
            title_para = doc.add_paragraph(f"{run_id}．{title}")
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in title_para.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(12)

            for header, field in zip(RUNWAY_HEADERS_ZH, RUNWAY_FIELD_ORDER):
                if field in ("id", "runway_name"):
                    continue
                value = row.get(field, "")
                para = doc.add_paragraph(style="List Bullet")
                bullet = para.add_run(f"{header}：{value}")
                bullet.font.name = "Times New Roman"
                bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                bullet.font.size = Pt(11)

        doc.save(filename)

def _to_str(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, Decimal):
        return str(value)
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d %H:%M:%S")
    if isinstance(value, (bytes, bytearray)):
        try:
            return base64.b64encode(value).decode("ascii")
        except Exception:
            return "[binary]"
    return str(value)


def runway_to_dict(obj: Any) -> Dict[str, str]:
    return {field: _to_str(getattr(obj, field, None)) for field in RUNWAY_FIELD_ORDER}
