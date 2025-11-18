from __future__ import annotations

import base64
import os
from datetime import datetime
from decimal import Decimal
from typing import Any, Dict, List, Optional, Sequence

from PyQt6.QtCore import QThread, pyqtSignal
from PyQt6.QtWidgets import (
    QComboBox,
    QDialog,
    QFileDialog,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QMessageBox,
    QPushButton,
    QProgressBar,
    QVBoxLayout,
)

SHELTER_FIELD_ORDER: Sequence[str] = (
    "id",
    "shelter_code",
    "shelter_name",
    "country",
    "base",
    "shelter_picture",
    "shelter_length",
    "shelter_width",
    "shelter_height",
    "cave_width",
    "cave_height",
    "structural_form",
    "door_material",
    "door_thick",
    "mask_layer_material",
    "mask_layer_thick",
    "soil_layer_material",
    "soil_layer_thick",
    "disper_layer_material",
    "disper_layer_thick",
    "disper_layer_reinforcement",
    "structure_layer_material",
    "structure_layer_thick",
    "structure_layer_reinforcement",
    "explosion_resistance",
    "anti_kinetic",
    "resistance_depth",
    "nuclear_blast",
    "radiation_shielding",
    "fire_resistance",
    "shelter_status",
    "created_time",
    "updated_time"
)



SHELTER_HEADERS: Sequence[str] = (
    "编号",
    "掩蔽库代码",
    "掩蔽库名称",
    "国家/地区",
    "基地/部队",
    "掩蔽库照片",
    "库容净宽",
    "库容净高",
    "库容净长",
    "洞门宽度",
    "洞门高度",
    "结构形式",
    "门体材料",
    "门体厚度",
    "伪装层材料",
    "伪装层厚度",
    "遮弹层材料",
    "遮弹层厚度",
    "分散层材料",
    "分散层厚度",
    "分散层钢筋配置",
    "结构层材料",
    "结构层厚度",
    "结构层钢筋配置",
    "抗爆能力",
    "抗动能穿透",
    "抗穿透深度",
    "抗核冲波超压",
    "抗辐射屏蔽",
    "耐火极限",
    "掩蔽库状态",
    "创建时间",
    "更新时间"
)


class Target_Shelter_ExportWindow(QDialog):
    def __init__(self, parent, session_scope) -> None:
        super().__init__(parent)
        self.setWindowTitle("导出单机掩蔽库数据模型")
        self.setModal(True)
        self.resize(500, 180)

        self.session_scope = session_scope
        self.worker: Optional[ExportWorker] = None

        layout = QVBoxLayout(self)

        row1 = QHBoxLayout()
        row1.addWidget(QLabel("导出目录："))
        self.ed_dir = QLineEdit()
        self.btn_browse = QPushButton("浏览...")
        self.btn_browse.clicked.connect(self._on_browse)
        row1.addWidget(self.ed_dir, 1)
        row1.addWidget(self.btn_browse)
        layout.addLayout(row1)

        row2 = QHBoxLayout()
        row2.addWidget(QLabel("导出格式："))
        self.cmb_fmt = QComboBox()
        self.cmb_fmt.addItems(["Excel", "Word"])
        row2.addWidget(self.cmb_fmt, 1)
        layout.addLayout(row2)

        row3 = QHBoxLayout()
        row3.addStretch(1)
        self.btn_export = QPushButton("开始导出")
        self.btn_export.clicked.connect(self._on_export)
        row3.addWidget(self.btn_export)
        layout.addLayout(row3)

        self.progress = QProgressBar()
        self.progress.setRange(0, 100)
        layout.addWidget(self.progress)

        self.lbl_status = QLabel("")
        self.lbl_status.setStyleSheet("color: gray;")
        layout.addWidget(self.lbl_status)

    def _on_browse(self) -> None:
        directory = QFileDialog.getExistingDirectory(self, "选择导出目录", os.path.expanduser("~"))
        if directory:
            self.ed_dir.setText(directory)

    def _on_export(self) -> None:
        out_dir = self.ed_dir.text().strip()
        if not out_dir:
            QMessageBox.warning(self, "提示", "请先选择导出目录")
            return
        if not os.path.isdir(out_dir):
            QMessageBox.warning(self, "提示", "导出目录无效")
            return

        fmt = self.cmb_fmt.currentText().strip().lower()

        self.progress.setValue(0)
        self.lbl_status.setText("正在启动导出 ...")
        self.btn_export.setEnabled(False)

        self.worker = ExportWorker(self.session_scope, out_dir, fmt, self)
        self.worker.progress.connect(self.progress.setValue)
        self.worker.message.connect(self.lbl_status.setText)
        self.worker.error.connect(self._on_error)
        self.worker.done.connect(self._on_done)
        self.worker.finished.connect(lambda: self.btn_export.setEnabled(True))
        self.worker.start()

    def _on_error(self, message: str) -> None:
        self.lbl_status.setText(message)
        QMessageBox.critical(self, "导出失败", message)

    def _on_done(self, filename: str) -> None:
        self.lbl_status.setText(f"导出完成：{filename}")
        QMessageBox.information(self, "导出完成", f"已导出到：\n{filename}")




class ExportWorker(QThread):
    progress = pyqtSignal(int)
    message = pyqtSignal(str)
    error = pyqtSignal(str)
    done = pyqtSignal(str)

    def __init__(self, session_scope, out_dir: str, fmt: str, parent=None):
        super().__init__(parent)
        self.session_scope = session_scope
        self.out_dir = out_dir
        self.fmt = fmt.lower().strip()

    def run(self) -> None:  # pragma: no cover - involves GUI thread
        try:
            from target_model.sql_repository import SQLRepository
            from target_model.entities import AircraftShelter
        except Exception as exc:
            self.error.emit(f"导出依赖加载失败：{exc}")
            return

        try:
            self.message.emit("正在读取掩蔽库数据 ...")
            with self.session_scope() as session:
                repo = SQLRepository(session)
                shelters = repo.list_all(AircraftShelter)
        except Exception as exc:
            self.error.emit(f"读取掩蔽库数据失败：{exc}")
            return

        total = max(len(shelters), 1)
        self.progress.emit(10)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        if self.fmt == "excel":
            output = os.path.join(self.out_dir, f"Shelters_{timestamp}.xlsx")
        elif self.fmt == "word":
            output = os.path.join(self.out_dir, f"Shelters_{timestamp}.docx")
        else:
            self.error.emit(f"未知导出格式：{self.fmt}")
            return

        self.message.emit("正在整理数据 ...")
        rows: List[Dict[str, str]] = []
        for index, shelter in enumerate(shelters, start=1):
            rows.append(shelter_to_dict(shelter))
            percent = 10 + int(60 * index / total)
            self.progress.emit(min(percent, 70))

        try:
            if self.fmt == "excel":
                self._write_excel(output, rows)
            else:
                self._write_word(output, rows)
        except ImportError as exc:
            self.error.emit(
                f"缺少导出依赖：{exc}。Excel 导出需要 pandas+openpyxl，Word 导出需要 python-docx。"
            )
            return
        except Exception as exc:
            self.error.emit(f"写文件失败：{exc}")
            return

        self.progress.emit(100)
        self.message.emit("导出完成")
        self.done.emit(output)

    def _write_excel(self, filename: str, rows: List[Dict[str, str]]) -> None:
        self.message.emit("正在写入 Excel ...")
        import pandas as pd  # type: ignore
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side  # type: ignore
        from openpyxl.utils import get_column_letter  # type: ignore

        ordered = [
            {header: row.get(field, "") for field, header in zip(SHELTER_FIELD_ORDER, SHELTER_HEADERS)}
            for row in rows
        ]
        df = pd.DataFrame(ordered, columns=SHELTER_HEADERS)

        with pd.ExcelWriter(filename, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="Shelters")
            ws = writer.sheets["Shelters"]

            header_font = Font(bold=True)
            header_fill = PatternFill("solid", fgColor="DDDDDD")
            thin = Side(border_style="thin", color="999999")
            border = Border(top=thin, bottom=thin, left=thin, right=thin)
            center = Alignment(vertical="center")

            for row in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=ws.max_column):
                for cell in row:
                    cell.border = border
                    if cell.row == 1:
                        cell.font = header_font
                        cell.fill = header_fill
                        cell.alignment = center

            for col_idx in range(1, ws.max_column + 1):
                column = get_column_letter(col_idx)
                max_length = max(
                    (len(str(ws.cell(row=row, column=col_idx).value or "")) for row in range(1, ws.max_row + 1)),
                    default=0,
                )
                ws.column_dimensions[column].width = min(max(max_length + 2, 12), 60)

            ws.freeze_panes = "A2"

    def _write_word(self, filename: str, rows: List[Dict[str, str]]) -> None:
        self.message.emit("正在写入 Word ...")
        from docx import Document  # type: ignore
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.shared import Pt  # type: ignore

        document = Document()
        document.add_heading("Aircraft Shelters", level=1)


        for row in rows or [{}]:
            name = row.get("shelter_name", "")
            code = row.get("shelter_code", "")
            shelter_id = row.get("id", "")
            if name and code:
                display_title = f"{code} ： {name}"
            else:
                display_title = name or code or "Unnamed Shelter"
            heading_parts = [part for part in (str(shelter_id).strip(), display_title) if part]
            title_paragraph = document.add_paragraph(" ： ".join(heading_parts))
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in title_paragraph.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "??")
                run.font.size = Pt(12)

            for header, field in zip(SHELTER_HEADERS, SHELTER_FIELD_ORDER):
                if field in ("id", "shelter_code", "shelter_name"):
                    continue
                value = row.get(field, "")
                paragraph = document.add_paragraph(style="List Bullet")
                run = paragraph.add_run(f"{header}: {value}")
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "??")
                run.font.size = Pt(11)

        document.save(filename)


def _to_str(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, Decimal):
        return str(value)
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d %H:%M:%S")
    if isinstance(value, (bytes, bytearray)):
        try:
            return base64.b64encode(value).decode("ascii")
        except Exception:
            return "[binary]"
    return str(value)


def shelter_to_dict(obj: Any) -> Dict[str, str]:
    return {field: _to_str(getattr(obj, field, None)) for field in SHELTER_FIELD_ORDER}
