"""
毁伤评估报告导出模块
支持导出为PDF和Word格式
"""
import os
import sys
from typing import Tuple, Optional, Dict, Any

# 添加项目根目录到路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from loguru import logger
from DBCode.DBHelper import DBHelper
from damage_models.sql_repository_dbhelper import (
    AssessmentReportRepository,
    AssessmentResultRepository,
    DamageSceneRepository,
    DamageParameterRepository
)


def get_report_full_data(report_id: int) -> Optional[Dict[str, Any]]:
    """
    根据报告ID获取完整的报告数据,包括所有关联表的详细信息

    返回字典结构:
    {
        'report': AssessmentReport对象,
        'result': AssessmentResult对象,
        'scene': DamageScene对象,
        'parameter': DamageParameter对象,
        'ammunition': Ammunition对象,
        'target': Target对象(根据类型可能是Runway/Shelter/UCC),
        'target_type_name': str (目标类型名称)
    }
    """
    db = DBHelper()
    try:
        # 1. 获取报告基本信息
        report_repo = AssessmentReportRepository(db)
        report = report_repo.get_by_id(report_id)

        if not report:
            logger.error(f"未找到ID为{report_id}的报告")
            return None

        data = {'report': report}

        # 2. 获取毁伤结果
        if report.DAID:
            result_repo = AssessmentResultRepository(db)
            result = result_repo.get_by_id(report.DAID)
            data['result'] = result
        else:
            data['result'] = None

        # 3. 获取毁伤场景
        if report.DSID:
            scene_repo = DamageSceneRepository(db)
            scene = scene_repo.get_by_id(report.DSID)
            data['scene'] = scene
        else:
            data['scene'] = None

        # 4. 获取毁伤参数
        if report.DPID:
            param_repo = DamageParameterRepository(db)
            parameter = param_repo.get_by_id(report.DPID)
            data['parameter'] = parameter
        else:
            data['parameter'] = None

        # 5. 获取弹药信息
        if report.AMID:
            sql = "SELECT * FROM Ammunition_Info WHERE AMID = %s"
            ammo_rows = db.execute_query(sql, [report.AMID])
            if ammo_rows:
                # 创建一个简单对象来存储弹药数据
                class Ammunition:
                    def __init__(self, row_dict):
                        for key, value in row_dict.items():
                            setattr(self, key, value)
                ammunition = Ammunition(ammo_rows[0])
                data['ammunition'] = ammunition
            else:
                data['ammunition'] = None
        else:
            data['ammunition'] = None

        # 6. 根据目标类型获取目标信息
        target_type_map = {
            1: '机场跑道',
            2: '单机掩蔽库',
            3: '地下指挥所'
        }

        data['target_type_name'] = target_type_map.get(report.TargetType, '未知')

        if report.TargetID and report.TargetType:
            target = None

            class Target:
                def __init__(self, row_dict):
                    for key, value in row_dict.items():
                        setattr(self, key, value)

            if report.TargetType == 1:  # 机场跑道
                sql = "SELECT * FROM Runway_Info WHERE RunwayID = %s"
                target_rows = db.execute_query(sql, [report.TargetID])
                if target_rows:
                    target = Target(target_rows[0])
            elif report.TargetType == 2:  # 单机掩蔽库
                sql = "SELECT * FROM Shelter_Info WHERE ShelterID = %s"
                target_rows = db.execute_query(sql, [report.TargetID])
                if target_rows:
                    target = Target(target_rows[0])
            elif report.TargetType == 3:  # 地下指挥所
                sql = "SELECT * FROM UCC_Info WHERE UCCID = %s"
                target_rows = db.execute_query(sql, [report.TargetID])
                if target_rows:
                    target = Target(target_rows[0])
            data['target'] = target
        else:
            data['target'] = None

        return data

    except Exception as e:
        logger.exception(f"获取报告完整数据失败: {e}")
        return None
    finally:
        db.close()


def export_report_to_pdf(report_id: int, output_path: str) -> Tuple[bool, str]:
    """
    导出报告为PDF格式

    Args:
        report_id: 报告ID
        output_path: 输出文件路径

    Returns:
        (成功标志, 消息)
    """
    try:
        # 检查reportlab库
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError:
            return False, "需要安装reportlab库: pip install reportlab"

        # 获取完整报告数据
        data = get_report_full_data(report_id)
        if not data:
            return False, "无法获取报告数据"

        report = data['report']
        result = data.get('result')
        scene = data.get('scene')
        parameter = data.get('parameter')
        ammunition = data.get('ammunition')
        target = data.get('target')
        target_type_name = data.get('target_type_name', '未知')

        # 注册中文字体 - 使用reportlab内置的CJK字体支持（不依赖本机字体）
        font_name = 'Helvetica'  # 默认字体

        try:
            # 使用reportlab内置的CJK字体（无需本地字体文件）
            # STSong-Light 是 Adobe 的简体中文宋体，reportlab 原生支持
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont

            # 尝试多个CJK字体选项
            cjk_fonts = [
                'STSong-Light',      # 简体中文宋体
                'MSung-Light',       # 繁体中文明体
                'HeiseiMin-W3',      # 日文明朝体（也支持中文）
                'HeiseiKakuGo-W5',   # 日文黑体（也支持中文）
            ]

            for cjk_font in cjk_fonts:
                try:
                    pdfmetrics.registerFont(UnicodeCIDFont(cjk_font))
                    font_name = cjk_font
                    logger.info(f"成功注册CJK字体: {cjk_font}（无需本地字体文件）")
                    break
                except Exception as e:
                    logger.debug(f"CJK字体 {cjk_font} 注册失败: {e}")
                    continue

            # 如果CJK字体都失败，尝试系统字体作为备选
            if font_name == 'Helvetica':
                logger.warning("CJK字体注册失败，尝试使用系统字体")
                font_paths = [
                    r"C:\Windows\Fonts\simsun.ttc",
                    r"C:\Windows\Fonts\simsun.ttf",
                    "/usr/share/fonts/truetype/arphic/simsun.ttf",  # Linux
                    "/System/Library/Fonts/STHeiti Light.ttc",  # macOS
                ]

                for font_path in font_paths:
                    if os.path.exists(font_path):
                        try:
                            pdfmetrics.registerFont(TTFont('SimSun', font_path))
                            font_name = 'SimSun'
                            logger.info(f"成功注册系统字体: {font_path}")
                            break
                        except Exception as e:
                            logger.debug(f"系统字体 {font_path} 注册失败: {e}")
                            continue

            if font_name == 'Helvetica':
                logger.warning("未找到任何中文字体，PDF中的中文可能显示异常")

        except Exception as e:
            logger.error(f"字体注册失败: {e}, 使用默认字体")
            font_name = 'Helvetica'

        # 创建PDF文档
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        story = []

        # 设置样式
        styles = getSampleStyleSheet()

        # 主标题样式 - 不使用parent以确保字体设置生效
        title_style = ParagraphStyle(
            'CustomTitle',
            fontName=font_name,
            fontSize=18,
            leading=22,
            alignment=TA_CENTER,
            spaceAfter=30,
            textColor=colors.black
        )

        # 一级标题样式 - 不使用parent以确保字体设置生效
        h1_style = ParagraphStyle(
            'CustomH1',
            fontName=font_name,
            fontSize=14,
            leading=17,
            spaceAfter=12,
            spaceBefore=12,
            leftIndent=0,
            textColor=colors.black
        )

        # 正文样式 - 不使用parent以确保字体设置生效
        normal_style = ParagraphStyle(
            'CustomNormal',
            fontName=font_name,
            fontSize=10,
            leading=16,
            leftIndent=0,
            textColor=colors.black
        )

        # 添加标题
        story.append(Paragraph(f"{report.ReportName}", title_style))
        story.append(Spacer(1, 0.5*cm))

        # 一、评估报告概述
        story.append(Paragraph("一、评估报告概述", h1_style))
        story.append(Spacer(1, 0.3*cm))

        overview_data = [
            ['报告编号', report.ReportCode or ''],
            ['报告名称', report.ReportName or ''],
            ['报告生成日期', report.CreatedTime.strftime('%Y-%m-%d %H:%M:%S') if report.CreatedTime else '']
        ]

        overview_table = Table(overview_data, colWidths=[4*cm, 12*cm])
        overview_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(overview_table)
        story.append(Spacer(1, 0.5*cm))

        # 二、弹药毁伤数据模型
        story.append(Paragraph("二、弹药毁伤数据模型", h1_style))
        story.append(Spacer(1, 0.3*cm))

        if ammunition:
            ammo_data = [
                ['字段', '值'],
                ['弹药名称', getattr(ammunition, 'AMName', '') or ''],
                ['中文名称', getattr(ammunition, 'AMNameCN', '') or ''],
                ['弹药类型', getattr(ammunition, 'AMType', '') or ''],
                ['弹药型号', getattr(ammunition, 'AMModel', '') or ''],
                ['国家/地区', getattr(ammunition, 'Country', '') or ''],
                ['弹体全重(kg)', str(getattr(ammunition, 'AMWeight', '') or '')],
                ['战斗部类型', getattr(ammunition, 'WarheadType', '') or ''],
                ['战斗部名称', getattr(ammunition, 'WarheadName', '') or ''],
                ['装药量(kg)', str(getattr(ammunition, 'ChargeAmount', '') or '')],
                ['TNT当量(吨)', str(getattr(ammunition, 'TNTEquivalent', '') or '')],
                ['载机(投放平台)', getattr(ammunition, 'Carrier', '') or ''],
                ['制导方式', getattr(ammunition, 'GuidanceMode', '') or ''],
                ['投弹高度范围(m)', getattr(ammunition, 'DropHeight', '') or ''],
                ['投弹速度(km/h)', str(getattr(ammunition, 'DropSpeed', '') or '')],
                ['射程(km)', str(getattr(ammunition, 'FlightRange', '') or '')],
            ]

            # 战斗部特有参数 - 显示所有类型
            # 爆破战斗部参数
            ammo_data.append(['--- 爆破战斗部参数 ---', ''])
            ammo_data.append(['炸药成分', getattr(ammunition, 'EXBComponent', '') or ''])
            ammo_data.append(['炸药热爆(kJ/kg)', str(getattr(ammunition, 'EXBExplosion', '') or '')])
            ammo_data.append(['装药质量(kg)', str(getattr(ammunition, 'EXBWeight', '') or '')])

            # 聚能战斗部参数
            ammo_data.append(['--- 聚能战斗部参数 ---', ''])
            ammo_data.append(['炸药密度(g/cm³)', str(getattr(ammunition, 'EBDensity', '') or '')])
            ammo_data.append(['装药爆速(m/s)', str(getattr(ammunition, 'EBVelocity', '') or '')])
            ammo_data.append(['爆轰压(GPa)', str(getattr(ammunition, 'EBPressure', '') or '')])
            ammo_data.append(['药型罩材料', getattr(ammunition, 'EBCoverMaterial', '') or ''])
            ammo_data.append(['药型罩锥角(度)', str(getattr(ammunition, 'EBConeAngle', '') or '')])

            # 破片战斗部参数
            ammo_data.append(['--- 破片战斗部参数 ---', ''])
            ammo_data.append(['炸弹热爆(kJ/kg)', str(getattr(ammunition, 'FBBombExplosion', '') or '')])
            ammo_data.append(['破片形状', getattr(ammunition, 'FBFragmentShape', '') or ''])
            ammo_data.append(['破片表面积(mm²)', str(getattr(ammunition, 'FBSurfaceArea', '') or '')])
            ammo_data.append(['破片质量(g)', str(getattr(ammunition, 'FBFragmentWeight', '') or '')])
            ammo_data.append(['装药直径(mm)', str(getattr(ammunition, 'FBDiameter', '') or '')])
            ammo_data.append(['壳体质量(kg)', str(getattr(ammunition, 'FBShellWeight', '') or '')])

            # 穿甲战斗部参数
            ammo_data.append(['--- 穿甲战斗部参数 ---', ''])
            ammo_data.append(['弹丸质量(kg)', str(getattr(ammunition, 'ABBulletWeight', '') or '')])
            ammo_data.append(['弹丸直径(mm)', str(getattr(ammunition, 'ABDiameter', '') or '')])
            ammo_data.append(['弹丸头部长度(mm)', str(getattr(ammunition, 'ABHeadLength', '') or '')])

            # 子母弹战斗部参数
            ammo_data.append(['--- 子母弹战斗部参数 ---', ''])
            ammo_data.append(['母弹质量(kg)', str(getattr(ammunition, 'CBMBulletWeight', '') or '')])
            ammo_data.append(['母弹最大横截面(m²)', str(getattr(ammunition, 'CBMBulletSection', '') or '')])
            ammo_data.append(['母弹阻力系数', str(getattr(ammunition, 'CBMProjectile', '') or '')])
            ammo_data.append(['子弹数量', str(getattr(ammunition, 'CBSBulletCount', '') or '')])
            ammo_data.append(['子弹型号', getattr(ammunition, 'CBSBulletModel', '') or ''])
            ammo_data.append(['子弹质量(kg)', str(getattr(ammunition, 'CBSBulletWeight', '') or '')])
            ammo_data.append(['最大直径(mm)', str(getattr(ammunition, 'CBDiameter', '') or '')])
            ammo_data.append(['子弹参考长度(mm)', str(getattr(ammunition, 'CBSBulletLength', '') or '')])


            ammo_table = Table(ammo_data, colWidths=[5*cm, 11*cm])
            ammo_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), font_name),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ]))
            story.append(ammo_table)
        else:
            story.append(Paragraph("无弹药数据", normal_style))

        story.append(Spacer(1, 0.5*cm))

        # 三、打击目标数据模型
        story.append(Paragraph("三、打击目标数据模型", h1_style))
        story.append(Spacer(1, 0.3*cm))

        if target:
            target_data = [['字段', '值']]

            if report.TargetType == 1:  # 机场跑道
                target_data.extend([
                    ['目标类型', target_type_name],
                    ['跑道代码', getattr(target, 'RunwayCode', '') or ''],
                    ['跑道名称', getattr(target, 'RunwayName', '') or ''],
                    ['国家/地区', getattr(target, 'Country', '') or ''],
                    ['基地/部队', getattr(target, 'Base', '') or ''],
                    ['跑道长度(m)', str(getattr(target, 'RLength', '') or '')],
                    ['跑道宽度(m)', str(getattr(target, 'RWidth', '') or '')],
                    ['混凝土面层厚度(cm)', str(getattr(target, 'PCCSCThick', '') or '')],
                    ['水泥稳定碎石基层厚度(cm)', str(getattr(target, 'CTBCThick', '') or '')],
                    ['级配砂砾石垫层厚度(cm)', str(getattr(target, 'GCSSThick', '') or '')],
                    ['土基压实层厚度(cm)', str(getattr(target, 'CSThick', '') or '')],
                ])
            elif report.TargetType == 2:  # 单机掩蔽库
                target_data.extend([
                    ['目标类型', target_type_name],
                    ['掩蔽库代码', getattr(target, 'ShelterCode', '') or ''],
                    ['掩蔽库名称', getattr(target, 'ShelterName', '') or ''],
                    ['国家/地区', getattr(target, 'Country', '') or ''],
                    ['基地/部队', getattr(target, 'Base', '') or ''],
                    ['库容净宽(m)', str(getattr(target, 'ShelterWidth', '') or '')],
                    ['库容净高(m)', str(getattr(target, 'ShelterHeight', '') or '')],
                    ['库容净长(m)', str(getattr(target, 'ShelterLength', '') or '')],
                    ['结构形式', getattr(target, 'StructuralForm', '') or ''],
                    ['结构层材料', getattr(target, 'StructureLayerMaterial', '') or ''],
                    ['结构层厚度(cm)', str(getattr(target, 'StructureLayerThick', '') or '')],
                    ['防护层材料', getattr(target, 'MaskLayerMaterial', '') or ''],
                ])
            elif report.TargetType == 3:  # 地下指挥所
                target_data.extend([
                    ['目标类型', target_type_name],
                    ['指挥所代码', getattr(target, 'UCCCode', '') or ''],
                    ['指挥所名称', getattr(target, 'UCCName', '') or ''],
                    ['国家/地区', getattr(target, 'Country', '') or ''],
                    ['基地/部队', getattr(target, 'Base', '') or ''],
                    ['所在位置', getattr(target, 'Location', '') or ''],
                    ['岩层材料', getattr(target, 'RockLayerMaterials', '') or ''],
                    ['岩层厚度(m)', str(getattr(target, 'RockLayerThick', '') or '')],
                    ['防护层材料', getattr(target, 'ProtectiveLayerMaterial', '') or ''],
                    ['防护层厚度(m)', str(getattr(target, 'ProtectiveLayerThick', '') or '')],
                    ['指挥中心墙壁材料', getattr(target, 'UCCWallMaterials', '') or ''],
                    ['指挥中心墙壁厚度(m)', str(getattr(target, 'UCCWallThick', '') or '')],
                ])

            target_table = Table(target_data, colWidths=[5*cm, 11*cm])
            target_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), font_name),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ]))
            story.append(target_table)
        else:
            story.append(Paragraph("无目标数据", normal_style))

        story.append(Spacer(1, 0.5*cm))

        # 四、毁伤场景
        story.append(Paragraph("四、毁伤场景", h1_style))
        story.append(Spacer(1, 0.3*cm))

        if scene:
            scene_data = [
                ['字段', '值'],
                ['场景编号', scene.DSCode or ''],
                ['场景名称', scene.DSName or ''],
                ['进攻方', scene.DSOffensive or ''],
                ['假想敌', scene.DSDefensive or ''],
                ['所在战场', scene.DSBattle or ''],
            ]

            scene_table = Table(scene_data, colWidths=[4*cm, 12*cm])
            scene_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), font_name),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(scene_table)
        else:
            story.append(Paragraph("无场景数据", normal_style))

        story.append(Spacer(1, 0.5*cm))

        # 五、毁伤参数
        story.append(Paragraph("五、毁伤参数", h1_style))
        story.append(Spacer(1, 0.3*cm))

        if parameter:
            param_data = [
                ['字段', '值'],
                ['投放平台', parameter.Carrier or ''],
                ['制导方式', parameter.GuidanceMode or ''],
                ['战斗部类型', parameter.WarheadType or ''],
                ['装药量(kg)', str(parameter.ChargeAmount or '')],
                ['投弹高度', parameter.DropHeight or ''],
                ['投弹速度(m/s)', str(parameter.DropSpeed or '')],
                ['投弹方式', parameter.DropMode or ''],
                ['射程(km)', str(parameter.FlightRange or '')],
                ['电磁干扰等级', parameter.ElectroInterference or ''],
                ['天气状况', parameter.WeatherConditions or ''],
                ['环境风速(m/s)', str(parameter.WindSpeed or '')],
            ]

            param_table = Table(param_data, colWidths=[4*cm, 12*cm])
            param_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), font_name),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(param_table)
        else:
            story.append(Paragraph("无参数数据", normal_style))

        story.append(Spacer(1, 0.5*cm))

        # 六、毁伤能力计算
        story.append(Paragraph("六、毁伤能力计算", h1_style))
        story.append(Spacer(1, 0.3*cm))

        if result:
            result_data = [
                ['字段', '值'],
                ['弹坑深度(m)', str(result.DADepth or '')],
                ['弹坑直径(m)', str(result.DADiameter or '')],
                ['弹坑容积(m³)', str(result.DAVolume or '')],
                ['弹坑面积(m²)', str(result.DAArea or '')],
                ['弹坑长度(m)', str(result.DALength or '')],
                ['弹坑宽度(m)', str(result.DAWidth or '')],
                ['结构破坏程度', str(result.Discturction or '')],
            ]

            result_table = Table(result_data, colWidths=[4*cm, 12*cm])
            result_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), font_name),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(result_table)
        else:
            story.append(Paragraph("无毁伤计算结果", normal_style))

        story.append(Spacer(1, 0.5*cm))

        # 七、毁伤评估
        story.append(Paragraph("七、毁伤评估", h1_style))
        story.append(Spacer(1, 0.3*cm))

        # 毁伤等级
        degree_data = [['毁伤等级', report.DamageDegree or '']]
        degree_table = Table(degree_data, colWidths=[4*cm, 12*cm])
        degree_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(degree_table)
        story.append(Spacer(1, 0.2*cm))

        # 评估结论 - 支持换行
        story.append(Paragraph("评估结论:", ParagraphStyle(
            'CommentTitle',
            parent=normal_style,
            fontName=font_name,
            fontSize=10,
            textColor=colors.black,
            spaceAfter=6
        )))

        comment_text = report.Comment or '无'
        comment_text = comment_text.replace('\n', '<br/>')
        comment_para = Paragraph(comment_text, ParagraphStyle(
            'CommentContent',
            parent=normal_style,
            fontName=font_name,
            fontSize=10,
            leading=16,
            leftIndent=20,
            spaceAfter=10
        ))
        story.append(comment_para)
        story.append(Spacer(1, 0.5*cm))

        # 八、报告人员信息
        story.append(Paragraph("八、报告人员信息", h1_style))
        story.append(Spacer(1, 0.3*cm))

        staff_data = [
            ['报告操作人员ID', str(report.Creator or '')],
            ['报告审核人', report.Reviewer or ''],
        ]

        staff_table = Table(staff_data, colWidths=[4*cm, 12*cm])
        staff_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(staff_table)

        # 生成PDF
        doc.build(story)

        logger.info(f"报告已成功导出为PDF: {output_path}")
        return True, "导出成功"

    except Exception as e:
        logger.exception(f"导出PDF失败: {e}")
        return False, str(e)


def set_paragraph_font(paragraph, font_name='宋体'):
    """完整设置段落字体为宋体（包括所有字体属性）"""
    from docx.oxml.ns import qn
    for run in paragraph.runs:
        run.font.name = font_name
        # 设置所有字体属性确保生效
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 东亚文字
        run._element.rPr.rFonts.set(qn('w:ascii'), font_name)      # ASCII字符
        run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)      # 高位ANSI
        run._element.rPr.rFonts.set(qn('w:cs'), font_name)         # 复杂脚本


def set_cell_font(cell, font_name='宋体'):
    """设置表格单元格字体为宋体"""
    from docx.oxml.ns import qn
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def export_report_to_word(report_id: int, output_path: str) -> Tuple[bool, str]:
    """
    导出报告为Word格式

    Args:
        report_id: 报告ID
        output_path: 输出文件路径

    Returns:
        (成功标志, 消息)
    """
    try:
        # 检查python-docx库
        try:
            from docx import Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from docx.oxml.ns import qn
        except ImportError:
            return False, "需要安装python-docx库: pip install python-docx"

        # 获取完整报告数据
        data = get_report_full_data(report_id)
        if not data:
            return False, "无法获取报告数据"

        report = data['report']
        result = data.get('result')
        scene = data.get('scene')
        parameter = data.get('parameter')
        ammunition = data.get('ammunition')
        target = data.get('target')
        target_type_name = data.get('target_type_name', '未知')

        # 创建Word文档
        doc = Document()

        # 设置中文字体 - 宋体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 设置标题字体为宋体（包括Title和各级Heading）
        # Title样式（level 0）
        try:
            title_style = doc.styles['Title']
            title_style.font.name = '宋体'
            title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        except:
            pass

        # Heading 1-9 样式
        for level in range(1, 10):
            try:
                heading_style = doc.styles[f'Heading {level}']
                heading_style.font.name = '宋体'
                heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            except:
                pass

        # 标题
        title = doc.add_heading(report.ReportName, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 使用辅助函数完整设置字体
        set_paragraph_font(title, '宋体')

        # 一、评估报告概述
        h1_para = doc.add_heading('一、评估报告概述', 1)
        set_paragraph_font(h1_para, '宋体')
        doc.add_paragraph()

        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        table.rows[0].cells[0].text = '报告编号'
        table.rows[0].cells[1].text = report.ReportCode or ''
        table.rows[1].cells[0].text = '报告名称'
        table.rows[1].cells[1].text = report.ReportName or ''
        table.rows[2].cells[0].text = '报告生成日期'
        table.rows[2].cells[1].text = report.CreatedTime.strftime('%Y-%m-%d %H:%M:%S') if report.CreatedTime else ''

        # 设置表格字体为宋体
        for row in table.rows:
            for cell in row.cells:
                set_cell_font(cell)

        doc.add_paragraph()

        # 二、弹药毁伤数据模型
        h1_para = doc.add_heading('二、弹药毁伤数据模型', 1)
        set_paragraph_font(h1_para, '宋体')
        doc.add_paragraph()

        if ammunition:
            # 基本弹药信息
            basic_rows = [
                ['字段', '值'],
                ['弹药名称', getattr(ammunition, 'AMName', '') or ''],
                ['中文名称', getattr(ammunition, 'AMNameCN', '') or ''],
                ['弹药类型', getattr(ammunition, 'AMType', '') or ''],
                ['弹药型号', getattr(ammunition, 'AMModel', '') or ''],
                ['国家/地区', getattr(ammunition, 'Country', '') or ''],
                ['弹体全重(kg)', str(getattr(ammunition, 'AMWeight', '') or '')],
                ['战斗部类型', getattr(ammunition, 'WarheadType', '') or ''],
                ['战斗部名称', getattr(ammunition, 'WarheadName', '') or ''],
                ['装药量(kg)', str(getattr(ammunition, 'ChargeAmount', '') or '')],
                ['TNT当量(吨)', str(getattr(ammunition, 'TNTEquivalent', '') or '')],
                ['载机(投放平台)', getattr(ammunition, 'Carrier', '') or ''],
                ['制导方式', getattr(ammunition, 'GuidanceMode', '') or ''],
                ['投弹高度范围(m)', getattr(ammunition, 'DropHeight', '') or ''],
                ['投弹速度(km/h)', str(getattr(ammunition, 'DropSpeed', '') or '')],
                ['射程(km)', str(getattr(ammunition, 'FlightRange', '') or '')],
            ]

            # 战斗部特有参数 - 显示所有类型
            special_rows = []

            # 爆破战斗部参数
            special_rows.extend([
                ['--- 爆破战斗部参数 ---', ''],
                ['炸药成分', getattr(ammunition, 'EXBComponent', '') or ''],
                ['炸药热爆(kJ/kg)', str(getattr(ammunition, 'EXBExplosion', '') or '')],
                ['装药质量(kg)', str(getattr(ammunition, 'EXBWeight', '') or '')],
            ])

            # 聚能战斗部参数
            special_rows.extend([
                ['--- 聚能战斗部参数 ---', ''],
                ['炸药密度(g/cm³)', str(getattr(ammunition, 'EBDensity', '') or '')],
                ['装药爆速(m/s)', str(getattr(ammunition, 'EBVelocity', '') or '')],
                ['爆轰压(GPa)', str(getattr(ammunition, 'EBPressure', '') or '')],
                ['药型罩材料', getattr(ammunition, 'EBCoverMaterial', '') or ''],
                ['药型罩锥角(度)', str(getattr(ammunition, 'EBConeAngle', '') or '')],
            ])

            # 破片战斗部参数
            special_rows.extend([
                ['--- 破片战斗部参数 ---', ''],
                ['炸弹热爆(kJ/kg)', str(getattr(ammunition, 'FBBombExplosion', '') or '')],
                ['破片形状', getattr(ammunition, 'FBFragmentShape', '') or ''],
                ['破片表面积(mm²)', str(getattr(ammunition, 'FBSurfaceArea', '') or '')],
                ['破片质量(g)', str(getattr(ammunition, 'FBFragmentWeight', '') or '')],
                ['装药直径(mm)', str(getattr(ammunition, 'FBDiameter', '') or '')],
                ['壳体质量(kg)', str(getattr(ammunition, 'FBShellWeight', '') or '')],
            ])

            # 穿甲战斗部参数
            special_rows.extend([
                ['--- 穿甲战斗部参数 ---', ''],
                ['弹丸质量(kg)', str(getattr(ammunition, 'ABBulletWeight', '') or '')],
                ['弹丸直径(mm)', str(getattr(ammunition, 'ABDiameter', '') or '')],
                ['弹丸头部长度(mm)', str(getattr(ammunition, 'ABHeadLength', '') or '')],
            ])

            # 子母弹战斗部参数
            special_rows.extend([
                ['--- 子母弹战斗部参数 ---', ''],
                ['母弹质量(kg)', str(getattr(ammunition, 'CBMBulletWeight', '') or '')],
                ['母弹最大横截面(m²)', str(getattr(ammunition, 'CBMBulletSection', '') or '')],
                ['母弹阻力系数', str(getattr(ammunition, 'CBMProjectile', '') or '')],
                ['子弹数量', str(getattr(ammunition, 'CBSBulletCount', '') or '')],
                ['子弹型号', getattr(ammunition, 'CBSBulletModel', '') or ''],
                ['子弹质量(kg)', str(getattr(ammunition, 'CBSBulletWeight', '') or '')],
                ['最大直径(mm)', str(getattr(ammunition, 'CBDiameter', '') or '')],
                ['子弹参考长度(mm)', str(getattr(ammunition, 'CBSBulletLength', '') or '')],
            ])

            # 合并所有行并创建表格
            all_rows = basic_rows + special_rows
            table = doc.add_table(rows=len(all_rows), cols=2)
            table.style = 'Light Grid Accent 1'

            for idx, (field, value) in enumerate(all_rows):
                cell0 = table.rows[idx].cells[0]
                cell1 = table.rows[idx].cells[1]
                cell0.text = field
                cell1.text = value
                set_cell_font(cell0)
                set_cell_font(cell1)
        else:
            doc.add_paragraph('无弹药数据')

        doc.add_paragraph()

        # 三、打击目标数据模型
        h1_para = doc.add_heading('三、打击目标数据模型', 1)
        set_paragraph_font(h1_para, '宋体')
        doc.add_paragraph()

        if target:
            if report.TargetType == 1:  # 机场跑道
                table = doc.add_table(rows=12, cols=2)
                table.style = 'Light Grid Accent 1'
                table.rows[0].cells[0].text = '字段'
                table.rows[0].cells[1].text = '值'
                table.rows[1].cells[0].text = '目标类型'
                table.rows[1].cells[1].text = target_type_name
                table.rows[2].cells[0].text = '跑道代码'
                table.rows[2].cells[1].text = getattr(target, 'RunwayCode', '') or ''
                table.rows[3].cells[0].text = '跑道名称'
                table.rows[3].cells[1].text = getattr(target, 'RunwayName', '') or ''
                table.rows[4].cells[0].text = '国家/地区'
                table.rows[4].cells[1].text = getattr(target, 'Country', '') or ''
                table.rows[5].cells[0].text = '基地/部队'
                table.rows[5].cells[1].text = getattr(target, 'Base', '') or ''
                table.rows[6].cells[0].text = '跑道长度(m)'
                table.rows[6].cells[1].text = str(getattr(target, 'RLength', '') or '')
                table.rows[7].cells[0].text = '跑道宽度(m)'
                table.rows[7].cells[1].text = str(getattr(target, 'RWidth', '') or '')
                table.rows[8].cells[0].text = '混凝土面层厚度(cm)'
                table.rows[8].cells[1].text = str(getattr(target, 'PCCSCThick', '') or '')
                table.rows[9].cells[0].text = '水泥稳定碎石基层厚度(cm)'
                table.rows[9].cells[1].text = str(getattr(target, 'CTBCThick', '') or '')
                table.rows[10].cells[0].text = '级配砂砾石垫层厚度(cm)'
                table.rows[10].cells[1].text = str(getattr(target, 'GCSSThick', '') or '')
                table.rows[11].cells[0].text = '土基压实层厚度(cm)'
                table.rows[11].cells[1].text = str(getattr(target, 'CSThick', '') or '')

                # 设置表格字体为宋体
                for row in table.rows:
                    for cell in row.cells:
                        set_cell_font(cell)
            elif report.TargetType == 2:  # 单机掩蔽库
                table = doc.add_table(rows=13, cols=2)
                table.style = 'Light Grid Accent 1'
                table.rows[0].cells[0].text = '字段'
                table.rows[0].cells[1].text = '值'
                table.rows[1].cells[0].text = '目标类型'
                table.rows[1].cells[1].text = target_type_name
                table.rows[2].cells[0].text = '掩蔽库代码'
                table.rows[2].cells[1].text = getattr(target, 'ShelterCode', '') or ''
                table.rows[3].cells[0].text = '掩蔽库名称'
                table.rows[3].cells[1].text = getattr(target, 'ShelterName', '') or ''
                table.rows[4].cells[0].text = '国家/地区'
                table.rows[4].cells[1].text = getattr(target, 'Country', '') or ''
                table.rows[5].cells[0].text = '基地/部队'
                table.rows[5].cells[1].text = getattr(target, 'Base', '') or ''
                table.rows[6].cells[0].text = '库容净宽(m)'
                table.rows[6].cells[1].text = str(getattr(target, 'ShelterWidth', '') or '')
                table.rows[7].cells[0].text = '库容净高(m)'
                table.rows[7].cells[1].text = str(getattr(target, 'ShelterHeight', '') or '')
                table.rows[8].cells[0].text = '库容净长(m)'
                table.rows[8].cells[1].text = str(getattr(target, 'ShelterLength', '') or '')
                table.rows[9].cells[0].text = '结构形式'
                table.rows[9].cells[1].text = getattr(target, 'StructuralForm', '') or ''
                table.rows[10].cells[0].text = '结构层材料'
                table.rows[10].cells[1].text = getattr(target, 'StructureLayerMaterial', '') or ''
                table.rows[11].cells[0].text = '结构层厚度(cm)'
                table.rows[11].cells[1].text = str(getattr(target, 'StructureLayerThick', '') or '')
                table.rows[12].cells[0].text = '防护层材料'
                table.rows[12].cells[1].text = getattr(target, 'MaskLayerMaterial', '') or ''

                # 设置表格字体为宋体
                for row in table.rows:
                    for cell in row.cells:
                        set_cell_font(cell)
            elif report.TargetType == 3:  # 地下指挥所
                table = doc.add_table(rows=13, cols=2)
                table.style = 'Light Grid Accent 1'
                table.rows[0].cells[0].text = '字段'
                table.rows[0].cells[1].text = '值'
                table.rows[1].cells[0].text = '目标类型'
                table.rows[1].cells[1].text = target_type_name
                table.rows[2].cells[0].text = '指挥所代码'
                table.rows[2].cells[1].text = getattr(target, 'UCCCode', '') or ''
                table.rows[3].cells[0].text = '指挥所名称'
                table.rows[3].cells[1].text = getattr(target, 'UCCName', '') or ''
                table.rows[4].cells[0].text = '国家/地区'
                table.rows[4].cells[1].text = getattr(target, 'Country', '') or ''
                table.rows[5].cells[0].text = '基地/部队'
                table.rows[5].cells[1].text = getattr(target, 'Base', '') or ''
                table.rows[6].cells[0].text = '所在位置'
                table.rows[6].cells[1].text = getattr(target, 'Location', '') or ''
                table.rows[7].cells[0].text = '岩层材料'
                table.rows[7].cells[1].text = getattr(target, 'RockLayerMaterials', '') or ''
                table.rows[8].cells[0].text = '岩层厚度(m)'
                table.rows[8].cells[1].text = str(getattr(target, 'RockLayerThick', '') or '')
                table.rows[9].cells[0].text = '防护层材料'
                table.rows[9].cells[1].text = getattr(target, 'ProtectiveLayerMaterial', '') or ''
                table.rows[10].cells[0].text = '防护层厚度(m)'
                table.rows[10].cells[1].text = str(getattr(target, 'ProtectiveLayerThick', '') or '')
                table.rows[11].cells[0].text = '指挥中心墙壁材料'
                table.rows[11].cells[1].text = getattr(target, 'UCCWallMaterials', '') or ''
                table.rows[12].cells[0].text = '指挥中心墙壁厚度(m)'
                table.rows[12].cells[1].text = str(getattr(target, 'UCCWallThick', '') or '')

                # 设置表格字体为宋体
                for row in table.rows:
                    for cell in row.cells:
                        set_cell_font(cell)
        else:
            doc.add_paragraph('无目标数据')

        doc.add_paragraph()

        # 四、毁伤场景
        h1_para = doc.add_heading('四、毁伤场景', 1)
        set_paragraph_font(h1_para, '宋体')
        doc.add_paragraph()

        if scene:
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Light Grid Accent 1'
            table.rows[0].cells[0].text = '字段'
            table.rows[0].cells[1].text = '值'
            table.rows[1].cells[0].text = '场景编号'
            table.rows[1].cells[1].text = scene.DSCode or ''
            table.rows[2].cells[0].text = '场景名称'
            table.rows[2].cells[1].text = scene.DSName or ''
            table.rows[3].cells[0].text = '进攻方'
            table.rows[3].cells[1].text = scene.DSOffensive or ''
            table.rows[4].cells[0].text = '假想敌'
            table.rows[4].cells[1].text = scene.DSDefensive or ''
            table.rows[5].cells[0].text = '所在战场'
            table.rows[5].cells[1].text = scene.DSBattle or ''

            # 设置表格字体为宋体
            for row in table.rows:
                for cell in row.cells:
                    set_cell_font(cell)
        else:
            doc.add_paragraph('无场景数据')

        doc.add_paragraph()

        # 五、毁伤参数
        h1_para = doc.add_heading('五、毁伤参数', 1)
        set_paragraph_font(h1_para, '宋体')
        doc.add_paragraph()

        if parameter:
            table = doc.add_table(rows=13, cols=2)
            table.style = 'Light Grid Accent 1'
            table.rows[0].cells[0].text = '字段'
            table.rows[0].cells[1].text = '值'
            table.rows[1].cells[0].text = '投放平台'
            table.rows[1].cells[1].text = parameter.Carrier or ''
            table.rows[2].cells[0].text = '制导方式'
            table.rows[2].cells[1].text = parameter.GuidanceMode or ''
            table.rows[3].cells[0].text = '战斗部类型'
            table.rows[3].cells[1].text = parameter.WarheadType or ''
            table.rows[4].cells[0].text = '装药量(kg)'
            table.rows[4].cells[1].text = str(parameter.ChargeAmount or '')
            table.rows[5].cells[0].text = '投弹高度'
            table.rows[5].cells[1].text = str(parameter.DropHeight or '')
            table.rows[6].cells[0].text = '投弹速度(m/s)'
            table.rows[6].cells[1].text = str(parameter.DropSpeed or '')
            table.rows[7].cells[0].text = '投弹方式'
            table.rows[7].cells[1].text = parameter.DropMode or ''
            table.rows[8].cells[0].text = '射程(km)'
            table.rows[8].cells[1].text = str(parameter.FlightRange or '')
            table.rows[9].cells[0].text = '电磁干扰等级'
            table.rows[9].cells[1].text = parameter.ElectroInterference or ''
            table.rows[10].cells[0].text = '天气状况'
            table.rows[10].cells[1].text = parameter.WeatherConditions or ''
            table.rows[11].cells[0].text = '环境风速(m/s)'
            table.rows[11].cells[1].text = str(parameter.WindSpeed or '')

            # 设置表格字体为宋体
            for row in table.rows:
                for cell in row.cells:
                    set_cell_font(cell)
        else:
            doc.add_paragraph('无参数数据')

        doc.add_paragraph()

        # 六、毁伤能力计算
        h1_para = doc.add_heading('六、毁伤能力计算', 1)
        set_paragraph_font(h1_para, '宋体')
        doc.add_paragraph()

        if result:
            table = doc.add_table(rows=8, cols=2)
            table.style = 'Light Grid Accent 1'
            table.rows[0].cells[0].text = '字段'
            table.rows[0].cells[1].text = '值'
            table.rows[1].cells[0].text = '弹坑深度(m)'
            table.rows[1].cells[1].text = str(result.DADepth or '')
            table.rows[2].cells[0].text = '弹坑直径(m)'
            table.rows[2].cells[1].text = str(result.DADiameter or '')
            table.rows[3].cells[0].text = '弹坑容积(m³)'
            table.rows[3].cells[1].text = str(result.DAVolume or '')
            table.rows[4].cells[0].text = '弹坑面积(m²)'
            table.rows[4].cells[1].text = str(result.DAArea or '')
            table.rows[5].cells[0].text = '弹坑长度(m)'
            table.rows[5].cells[1].text = str(result.DALength or '')
            table.rows[6].cells[0].text = '弹坑宽度(m)'
            table.rows[6].cells[1].text = str(result.DAWidth or '')
            table.rows[7].cells[0].text = '结构破坏程度'
            table.rows[7].cells[1].text = str(result.Discturction or '')

            # 设置表格字体为宋体
            for row in table.rows:
                for cell in row.cells:
                    set_cell_font(cell)
        else:
            doc.add_paragraph('无毁伤计算结果')

        doc.add_paragraph()

        # 七、毁伤评估
        h1_para = doc.add_heading('七、毁伤评估', 1)
        set_paragraph_font(h1_para, '宋体')
        doc.add_paragraph()

        # 毁伤等级
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        table.rows[0].cells[0].text = '毁伤等级'
        table.rows[0].cells[1].text = report.DamageDegree or ''

        # 设置表格字体为宋体
        for row in table.rows:
            for cell in row.cells:
                set_cell_font(cell)

        doc.add_paragraph()

        # 评估结论 - 支持换行（使用add_heading确保字体生效）
        h3_para = doc.add_heading('评估结论:', level=3)
        # 使用辅助函数完整设置字体
        set_paragraph_font(h3_para, '宋体')
        comment_text = report.Comment or '无'
        doc.add_paragraph(comment_text)

        doc.add_paragraph()

        # 八、报告人员信息
        h1_para = doc.add_heading('八、报告人员信息', 1)
        set_paragraph_font(h1_para, '宋体')
        doc.add_paragraph()

        table = doc.add_table(rows=2, cols=2)
        table.style = 'Light Grid Accent 1'
        table.rows[0].cells[0].text = '报告操作人员ID'
        table.rows[0].cells[1].text = str(report.Creator or '')
        table.rows[1].cells[0].text = '报告审核人'
        table.rows[1].cells[1].text = report.Reviewer or ''

        # 设置表格字体为宋体
        for row in table.rows:
            for cell in row.cells:
                set_cell_font(cell)

        # 保存文档
        doc.save(output_path)

        logger.info(f"报告已成功导出为Word: {output_path}")
        return True, "导出成功"

    except Exception as e:
        logger.exception(f"导出Word失败: {e}")
        return False, str(e)


def export_report_to_file(report_id: int, output_path: str, format: str = "pdf") -> Tuple[bool, str]:
    """
    导出报告到文件

    Args:
        report_id: 报告ID
        output_path: 输出文件路径
        format: 导出格式 ("pdf" 或 "word")

    Returns:
        (成功标志, 消息)
    """
    format = format.lower().strip()

    if format == "pdf":
        return export_report_to_pdf(report_id, output_path)
    elif format in ["word", "docx"]:
        return export_report_to_word(report_id, output_path)
    else:
        return False, f"不支持的导出格式: {format}"

